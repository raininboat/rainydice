# -*- coding: utf-8 -*-
'''
       ___  ___   _____  ____  __
      / _ \/ _ | /  _/ |/ /\ \/ /
     / , _/ __ |_/ //    /  \  / 
    /_/|_/_/ |_/___/_/|_/   /_/  

    RainyDice 跑团投掷机器人服务 by RainyZhou
        跑团日志 log 记录，DOCX 文档实现
    
    Copyright (C) 2021  RainyZhou  
                        Email: thunderain_zhou@163.com

    This file is part of RainyDice.

    RainyDice is free software: you can redistribute it and/or modify
    it under the terms of the GNU Affero General Public License as
    published by the Free Software Foundation, either version 3 of the
    License, or (at your option) any later version.

    RainyDice is distributed in the hope that it will be useful,
    but WITHOUT ANY WARRANTY; without even the implied warranty of
    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
    GNU Affero General Public License for more details.

    You should have received a copy of the GNU Affero General Public License
    along with RainyDice.  If not, see <https://www.gnu.org/licenses/>

'''
# 需使用 python-docx 库
import docx
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from rainydice.msgesacpe import messageEscape
from rainydice.dice_command.constant_key import LOG

class logdict(object):
    def __init__(self,key,line):
        self.logline = self.defaultline()
        for i,v in zip(key,line):
            self.logline.__setattr__(i,v)
        self.name = messageEscape.htmlunescape(self.logline.User_Name)
        
    class defaultline(object):
        def __getattribute__(self, name):
            return name
class logdocx(object):
    def __init__(self):
        self = docx.Document()
        self.styles['Normal'].font.name = u"Times New Roman"
        self.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        self.core_properties.author = 'Rainy Dice'
        self.core_properties.title = 'DICE LOG 跑团记录'
        

    def logline(self,logdict,color):
        logtime = logdict['time']
        name = logdict['name']
        textlist = str.splitlines(logdict['text'])
        
        for text in textlist:
            logtext = '<{name}> {text}'.format(name=name,text=text)
            para = self.add_paragraph()
            run = para.add_run()
            run.font.color.rgb = RGBColor(0xc0,0xc0,0xc0)
            run.add_text(logtime)
            run = para.add_run()
            run.font.color.rgb = RGBColor(LOG.colorid_hex[color])
            run.add_text(logtext)

    def logsave(self,path):
        self.save(path)

class loghtml(object):
    def __init__(self,path):
        self.file = open(path+'log.html','w', encoding = 'utf-8',errors='ignore')
        self.file.write(LOG.htmlheader)

    def logline(self,logdict,colorid):
        htmltamplate='''\
<span style="color: #C0C0C0;font-size: 1.5">{time}</span>
<span style="color: {color};font-size: 1.5">&lt;{name}&gt; {text}</span>
<br/>
'''
        color = LOG.colorid_str[colorid]
        textlist = str.splitlines(logdict['text'])
        for text in textlist:
            self.file.write(htmltamplate.format(time=logdict['time'],color=color,name=logdict['name'],))

    def logsave(self):
        self.file.write(LOG.htmlend)
        self.file.close()
