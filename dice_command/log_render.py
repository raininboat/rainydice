# -*- coding: utf-8 -*-
'''
       ___  ___   _____  ____  __
      / _ \/ _ | /  _/ |/ /\ \/ /
     / , _/ __ |_/ //    /  \  / 
    /_/|_/_/ |_/___/_/|_/   /_/  

    RainyDice 跑团投掷机器人服务 by RainyZhou
        跑团日志 log 记录， log文件生成模块
    
    Copyright (C) 2021  RainyZhou  
                        Email: thunderain_zhou@163.com

    This file is part of RainyDice.

    RainyDice is free software: you can redistribute it and/or modify
    it under the terms of the GNU Affero General Public License as
    published by the Free Software Foundation, either version 3 of the
    License, or (at your option) any later version.

    RainyDice is distributed in the hope that it will be useful,
    but WITHOUT ANY WARRANTY; without even the implied warranty of
    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
    GNU Affero General Public License for more details.

    You should have received a copy of the GNU Affero General Public License
    along with RainyDice.  If not, see <https://www.gnu.org/licenses/>

'''
# 需使用 python-docx 库
import docx
from docx.shared import RGBColor
from docx.oxml.ns import qn
# from docx.shared import Pt
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from time import strftime , localtime

from rainydice.msgesacpe import messageEscape
from rainydice.dice_command.constant_key import LOG

class LogLine(object):
    def __init__(self,line):
        self.logline = {}
        for i,v in zip(LOG.key,line):
            self.logline[i] = v
        self.name = messageEscape.htmlunescape(self.logline['User_Name'])
        self.text = messageEscape.htmlunescape(self.logline['User_Text'])
        self.id = self.logline['User_ID']
        self.pf = self.logline['Platform']
        self.timestamp = self.logline['Log_Time']
        self.time_arr = localtime(self.timestamp)


class LogDocx(object):
    def __init__(self,path):
        self.logpath=path+'log.docx'
        self.doc = docx.Document()
        self.doc.styles['Normal'].font.name = u"Calibri"
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        self.doc.styles['Normal'].paragraph_format.space_after=0
        self.doc.styles['Normal'].paragraph_format.space_before=0
        self.doc.core_properties.author = 'Rainy Dice'
        self.doc.core_properties.title = 'DICE LOG 跑团记录'

    def logline(self,logline:LogLine,colorid:int):
        logtime = strftime('%H:%M:%S',logline.time_arr)
        name = logline.name
        textlist = str.splitlines(logline.text)
        
        for text in textlist:
            logtext = '<{name}> {text}'.format(name=name,text=text)
            para = self.doc.add_paragraph()
            run = para.add_run()
            run.font.color.rgb = RGBColor(0xc0,0xc0,0xc0)
            run.add_text(logtime)
            run = para.add_run()
            r,g,b = LOG.colorid_hex[colorid]
            run.font.color.rgb = RGBColor(r,g,b)
            run.add_text(logtext)

    def logsave(self):
        self.doc.save(self.logpath)

class LogHtml(object):
    def __init__(self,path):
        self.file = open(path+'log.html','w', encoding = 'utf-8',errors='ignore')
        self.file.write(LOG.html_header)

    def logline(self,logline:LogLine,colorid:int):
        color = LOG.colorid_str[colorid]
        textlist = str.splitlines(logline.text)
        logtime = strftime('%H:%M:%S',logline.time_arr)
        for text in textlist:
            self.file.write(LOG.html_tamplate.format(time=logtime,color=color,name=logline.name,text=text))

    def logsave(self):
        self.file.write(LOG.html_end)
        self.file.close()

class LogTxtRaw(object):
    def __init__(self,path):
        self.file = open(path+'raw.txt','w',encoding='utf-8',errors='ignore')
    
    def logline(self,thisline:LogLine,colorid=None):
        name = thisline.name
        userid = thisline.id
        text=thisline.text
        logtime = strftime('%Y-%m-%d %H:%M:%S',thisline.time_arr)
        self.file.write(LOG.txtraw_tamplate.format(name=name,User_ID=userid,Time=logtime,text=text))


    def logsave(self):
        self.file.close()

class LogCsv(object):
    def __init__(self,path):
        self.file = open(path+'log_form.csv', 'w', encoding = 'utf-8-sig',errors='ignore')

    def logline(self,thisline:LogLine,colorid=None):
        loglist=[]
        for i in LOG.key:
            # str.replace(log_dict[i],'"','""')
            loglist.append('"{0}"'.format(str.replace(thisline.logline[i].__str__(),'"','""')))
        self.file.write(','.join(loglist))

    def logsave(self):
        self.file.close()
